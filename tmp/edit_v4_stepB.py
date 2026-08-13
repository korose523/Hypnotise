"""Step B: apply advisor reframing (Sec.10) + C1/C2/C4 fixes on top of v5 (output of Step A).
Input:  paper_en_submission_v5.docx
Output: paper_en_submission_v5.docx (in place)
"""
import re
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = r"E:/universal_bci_hypnosis/paper_en_submission_v5.docx"
doc = docx.Document(PATH)

def set_run_text(p, text, bold=False):
    if hasattr(p, '_p'):
        p = p._p
    for r in p.findall(qn('w:r')):
        p.remove(r)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        b = OxmlElement('w:b'); rPr.append(b)
    r.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t); p.append(r)

def rep(snippet, new_text):
    for p in doc.paragraphs:
        if snippet in p.text:
            set_run_text(p, new_text)
            return True
    raise SystemExit(f"SNIPPET NOT FOUND: {snippet[:60]}")

def new_para(style=None):
    p = OxmlElement('w:p')
    if style:
        pPr = OxmlElement('w:pPr'); ps = OxmlElement('w:pStyle'); ps.set(qn('w:val'), style)
        pPr.append(ps); p.append(pPr)
    return p

# ---------- Reframing: Abstract (lead with assets + two warnings) ----------
rep("Cross-dataset generalization under weak label supervision remains an open problem",
"""Cross-dataset generalization of EEG-based classification under weak, proxy-derived labels remains an open problem for altered-states research. We present a reproducible eight-dataset alignment pipeline that maps eight heterogeneous EEG corpora (712,832 windows; 697,906 with valid labels) to a common 14-channel EPOC+ montage with 63-dimensional spectral features, and we recover the real 1-9 arousal self-assessments for MAHNOB-HCI from session.xml metadata. As a benchmark, Random Forest classifiers are trained on seven source domains and evaluated on the held-out target under both zero-shot and 20%-participant few-shot calibration. The benchmark exposes two concrete methodological pitfalls rather than a performance result: (i) per-class recall shows every target collapsing to a single majority class, and (ii) a within-dataset upper-bound experiment (Table 3) shows that six of eight proxy label sets sit at or below three-class chance even when trained and tested on the same dataset, so the cross-dataset failure is a label-validity problem rather than a transfer-method problem. Across the eight targets, zero-shot accuracy averages 36.85% (95% CI 34.40-39.30) and calibrated 43.76% (41.77-45.75), but balanced accuracy stays at 33.01-35.62% (Cohen's kappa <= 0.068), i.e. at chance. We position this work as a preliminary methodological study whose contribution is the reusable alignment pipeline, the recovered MAHNOB labels, a quantitative estimate of split-leakage inflation, and a transparently reported negative result.""")

# ---------- Reframing: Sec.1 purpose ----------
rep("The purpose of this study is therefore twofold",
"""The purpose of this study is therefore to construct a reproducible cross-dataset alignment benchmark for proxy-labelled EEG state classification and to report, transparently, the two methodological pitfalls it reveals: (i) label collapse, in which proxy labels resolve to a single majority class, and (ii) label invalidity, in which proxy labels carry little learnable signal even within a single dataset. Rather than claiming strong classification performance, we treat the benchmark and its diagnostic findings as the primary contribution.""")

# ---------- Reframing: Sec.4.2 lead (link to within-domain Table 3) ----------
rep("presents the main results across all eight target domains over 20 random seeds",
"""Table 4 presents the main results across all eight target domains over 20 random seeds. Three-class random chance is 33.3%. Read against the within-dataset ceiling in Table 3, the cross-dataset results show that where a proxy label is learnable in-dataset (ds006437, within BAcc 68.2%) its transfer still collapses to chance-level balanced accuracy, and where the label is near-random in-dataset (the remaining six datasets) no transfer can be expected.""")

# ---------- Reframing: Sec.5.1 interpretation ----------
rep("The central finding of this study is negative and, we argue, informative",
"""The central finding of this study is negative and, we argue, informative. The within-domain upper bound (Table 3) shows that six of eight proxy label sets yield balanced accuracy at or below three-class chance (33.3%) even in the easiest same-dataset setting, so their labels carry essentially no learnable class signal. Only ds006437 (within BAcc 68.2%) has labels that are learnable within-dataset, yet its cross-dataset transfer still collapses to chance-level balanced accuracy (33.5%). Raw cross-dataset accuracy above chance was obtained on five of eight targets, yet balanced accuracy never exceeded 35.62% and Cohen's kappa never exceeded 0.068, and every apparent success decomposes into majority-class prediction once per-class recall is examined.""")

# ---------- Reframing: Sec.8 empirical-result paragraph (link Table 3) ----------
rep("The empirical result is negative. Overall zero-shot accuracy of 36.85%",
"""The empirical result is negative. Overall zero-shot accuracy of 36.85% sits 3.55pp above three-class chance, calibration adds 6.91pp in aggregate but that gain is confined to one target, and balanced accuracy never departs meaningfully from chance on any target. Read against the within-domain ceiling (Table 3), this is expected: six of eight proxy label sets are at or below chance even within-dataset, so no cross-dataset transfer could be expected for them, and the one learnable label set (ds006437) still fails to transfer.""")

# ---------- C1: Sec.3.2 Split unit — add GroupShuffleSplit sentence + C2 ds006437 separation ----------
rep("An 80/20 calibration/test split at the participant level guarantees that no participant contributes to both partitions",
"""Partitioning uses real participant identifiers throughout, derived as follows: MAHNOB participants from the <subject id> field of session.xml (27 participants); SEED participants from file-name numbers (10 participants present in the processed data); SEED_IV participants from the participant identifiers embedded in the feature filenames (15 participants); DREAMER, DEAP, FACED and ds004572 from their native participant identifiers; ds006437 is partitioned by session identifier (36 sessions from 9 participants) rather than by participant, because the public data does not expose a stable per-participant key. An 80/20 calibration/test split at the participant (or, for ds006437, session) level guarantees that no unit contributes to both partitions, eliminating the within-participant leakage present in trial-level partitioning. Partitions were generated with sklearn.model_selection.GroupShuffleSplit (n_splits = 1, test_size = 0.2), the unit identifier supplied as the groups argument and the experiment seed passed as random_state; the same function and the 20 config seeds generate both the calibration/test split and the within-domain upper-bound split (Table 3).""")

# ---------- C2(b): Table 1 ds006437 participants + footnote ----------
t0 = doc.tables[0]
for r in t0.rows:
    if r.cells[0].text.strip() == "ds006437":
        set_run_text(r.cells[0].paragraphs[0], "9 (36 sessions)++")
        break
# footnote paragraph after Table 1 caption
foot = None
for p in doc.paragraphs:
    if p.text.strip().startswith("Only 10 of the 15 publicly documented SEED"):
        foot = p
if foot is not None:
    fp = new_para()
    set_run_text(fp, "++ ds006437 comprises 9 participants across 36 sessions; it is partitioned by session identifier, so within-participant leakage cannot be fully excluded and its within-domain estimate (Table 3) uses session-level grouping.")
    foot._p.addnext(fp)

# ---------- C2(b): limitations table (Table 12 = doc.tables[11] after Step A insertion) add ds006437 session-leakage row ----------
lim = doc.tables[11]
new_cells = lim.add_row().cells
set_run_text(new_cells[0].paragraphs[0], "ds006437 is partitioned by session rather than participant identifier (9 participants, 36 sessions), so within-participant leakage cannot be fully excluded.")
set_run_text(new_cells[1].paragraphs[0], "Documented; noted wherever ds006437 is interpreted (Tables 3, 4, 8)")

# ---------- C4: funding / competing interests after Sec.7.3 ----------
body73 = None
for p in doc.paragraphs:
    if p.text.strip().startswith("A completed checklist mapping each STROBE"):
        body73 = p
if body73 is not None:
    p1 = new_para(); set_run_text(p1, "The authors received no specific funding for this work.")
    p2 = new_para(); set_run_text(p2, "The authors have declared that no competing interests exist.")
    body73._p.addnext(p2)
    body73._p.addnext(p1)

doc.save(PATH)
print("SAVED", PATH)
