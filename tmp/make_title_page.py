#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build PLOS ONE-compliant title_page.docx per the PLOS ONE
title/authors/affiliations formatting sample.

Rules applied:
- Title in sentence case (only first word + proper nouns capitalized).
- Author byline: names exactly as submitted, affiliation numbers and the
  '*' corresponding-author symbol in superscript; NO titles (Dr./Prof.).
- Affiliations listed small->large (Dept, Institution, City, Country),
  no ZIP / street / building numbers, no abbreviations, no positions.
- Corresponding author: email only + initials in parentheses (no address).
- ORCID iDs noted (Jung Minpo's is a required placeholder).
- CRediT contributions, Funding, Competing Interests sections.
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:/Users/Administrator/Downloads/title_page.docx"

doc = Document()
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
if doc.paragraphs and doc.paragraphs[0].text == '':
    p = doc.paragraphs[0]._element
    p.getparent().remove(p)

def line(text=None, bold=False, size=12, space_after=6, runs=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(space_after)
    if runs:
        for txt, kw in runs:
            r = p.add_run(txt)
            r.bold = kw.get('bold', False)
            if kw.get('sup'):
                r.font.superscript = True
            if kw.get('size'):
                r.font.size = Pt(kw['size'])
    elif text is not None:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
    return p

# ---- Title (sentence case, 16pt bold) ----
TITLE = ("Multi-source domain generalization with few-shot calibration for "
         "cross-dataset EEG hypnosis depth classification under proxy labels")
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = tp.add_run(TITLE)
tr.bold = True
tr.font.size = Pt(16)
tp.paragraph_format.space_after = Pt(6)

# ---- Short title (running head, <=100 chars per PLOS) ----
line("Short title: EEG hypnosis depth classification under proxy labels",
     bold=False, size=11, space_after=12)

# ---- Author byline (superscript numbers + * corresponding symbol) ----
bp = doc.add_paragraph()
bp.paragraph_format.line_spacing = 1.0
bp.paragraph_format.space_after = Pt(6)
for txt, sup in [("Weng Zexiao", False), ("1", True),
                 (", Jung Minpo", False), ("1", True), ("*", True)]:
    r = bp.add_run(txt)
    if sup:
        r.font.superscript = True

# ---- ORCID iDs ----
line("ORCID iDs:", bold=True, space_after=2)
line("Weng Zexiao: https://orcid.org/0009-0009-8600-8954", space_after=2)
line("Jung Minpo: [ORCID to be added \u2014 PLOS ONE requires an ORCID iD for the corresponding author]",
     space_after=8)

# ---- Affiliations (numbered, small->large, no titles/zip) ----
line("1 Department of Computer Engineering, Youngsan University, Yangsan, Republic of Korea",
     space_after=8)

# ---- Corresponding author (email only + initials) ----
line("* Corresponding author", space_after=2)
line("E-mail: minpo@ysu.ac.kr (JM)", space_after=8)

# ---- Author Contributions (CRediT) ----
line("Authors' Contributions (CRediT):", bold=True, space_after=2)
line("Weng Zexiao: Conceptualization, Data curation, Formal analysis, Investigation, "
     "Methodology, Resources, Software, Validation, Visualization, Writing \u2013 original "
     "draft, Writing \u2013 review & editing.", space_after=2)
line("Jung Minpo: Supervision, Writing \u2013 review & editing.", space_after=8)

# ---- Funding ----
line("Funding:", bold=True, space_after=2)
line("The authors received no specific funding for this work.", space_after=8)

# ---- Competing Interests ----
line("Competing Interests:", bold=True, space_after=2)
line("The authors have declared that no competing interests exist.", space_after=6)

# ---------- page numbers in footer + continuous line numbers ----------
footer = doc.sections[0].footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fld = OxmlElement('w:fldSimple')
fld.set(qn('w:instr'), ' PAGE ')
rr = OxmlElement('w:r')
tt = OxmlElement('w:t')
tt.text = '1'
rr.append(tt)
fld.append(rr)
fp._p.append(fld)

def add_line_numbers(document, count_by=1, restart='continuous', distance='360'):
    for sec in document.sections:
        sectPr = sec._sectPr
        ln = sectPr.find(qn('w:lnNumType'))
        if ln is None:
            ln = OxmlElement('w:lnNumType')
            sectPr.append(ln)
        ln.set(qn('w:countBy'), str(count_by))
        ln.set(qn('w:restart'), restart)
        ln.set(qn('w:distance'), distance)

add_line_numbers(doc)

doc.save(OUT)
print(f"Saved {OUT}")
