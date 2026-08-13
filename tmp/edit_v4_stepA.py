"""Step A: insert within-domain upper-bound section before old 4.1 in v4 docx,
and cascade-renumber 4.x subsections and Table 3..12 -> 4..13.
Base file: paper_en_submission_v4.docx (the advisor-approved version).
"""
import re, copy
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:/Users/Administrator/xwechat_files/wengzexiao_e4a2/msg/file/2026-08/3) paper_en_submission_v4.docx"
OUT = r"E:/universal_bci_hypnosis/paper_en_submission_v5.docx"

doc = docx.Document(SRC)

def set_run_text(p, text, bold=False):
    # accept either a raw w:p element or a docx Paragraph wrapper
    if hasattr(p, '_p'):
        p = p._p
    # clear existing runs, add one run
    for r in p.findall(qn('w:r')):
        p.remove(r)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        b = OxmlElement('w:b'); rPr.append(b)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)

def new_para(style=None):
    p = OxmlElement('w:p')
    if style:
        pPr = OxmlElement('w:pPr')
        ps = OxmlElement('w:pStyle'); ps.set(qn('w:val'), style)
        pPr.append(ps); p.append(pPr)
    return p

# ---------- 1. Shift Table N (3..12) -> N+1 in ALL paragraphs ----------
tbl_pat = re.compile(r'Table (\d+)')
def shift_table(m):
    n = int(m.group(1))
    if 3 <= n <= 12:
        return f"Table {n+1}"
    return m.group(0)
for p in doc.paragraphs:
    if tbl_pat.search(p.text):
        p.text = tbl_pat.sub(shift_table, p.text)

# ---------- 2. Shift headings "4.N " -> "4.(N+1) " and "Section 4.N" -> "Section 4.(N+1)" ----------
head_pat = re.compile(r'^(4\.\d+)\s')
def shift_head(m):
    num = m.group(1)
    major, minor = num.split('.')
    return f"{major}.{int(minor)+1} "
for p in doc.paragraphs:
    if p.style.name.startswith("Heading") and head_pat.match(p.text):
        p.text = head_pat.sub(shift_head, p.text)

sec_pat = re.compile(r'Section 4\.(\d+)')
def shift_sec(m):
    return f"Section 4.{int(m.group(1))+1}"
for p in doc.paragraphs:
    if sec_pat.search(p.text):
        p.text = sec_pat.sub(shift_sec, p.text)

# ---------- 3. Find anchor = old 4.1 now renamed to "4.2 Multi-source LODO performance" ----------
anchor = None
for p in doc.paragraphs:
    if p.text.startswith("4.2 Multi-source LODO performance"):
        anchor = p
        break
assert anchor is not None, "anchor 4.2 heading not found"

# ---------- 4. Build new content: intro para, heading, table ----------
intro = new_para()
set_run_text(intro,
    "Before reporting cross-dataset transfer we establish, for each dataset, the within-dataset "
    "classification ceiling: the same Random Forest is trained and tested on the same dataset under a "
    "participant-disjoint 80/20 split (sklearn GroupShuffleSplit, group = participant identifier; for "
    "ds006437 the only available unit is the session, because the public data carries 9 participants "
    "across 36 sessions, so its within-domain estimate uses session-level grouping). Identical 63-dimensional "
    "features, RF configuration (Table 2) and 20 seeds are used, so the within-domain and cross-dataset "
    "figures are directly comparable. We pre-specify the interpretation rule: a within-dataset balanced "
    "accuracy (BAcc) >= 40% indicates a learnable label signal, 35-40% is marginal, and < 35% means the "
    "proxy label is effectively noise even in the easiest (same-dataset) setting.")

head = new_para("Heading 2")
set_run_text(head, "4.1 Within-domain upper bound")

# table
cols = ["Dataset", "Within Acc (%)", "Within BAcc (%)", "Within kappa", "Cross ZS BAcc (%)", "Delta (Within - Cross)"]
rows = [
    ["DEAP",      "44.28 +/- 11.65", "35.72 +/- 9.17", "0.006 +/- 0.142", "35.62", "+0.10"],
    ["DREAMER",   "47.61 +/- 3.12",  "32.55 +/- 3.06", "0.007 +/- 0.052", "33.01", "-0.46"],
    ["MAHNOB",    "33.76 +/- 4.19",  "34.17 +/- 1.69", "0.008 +/- 0.025", "33.55", "+0.62"],
    ["SEED",      "33.56 +/- 0.43",  "33.33 +/- 0.00", "0.000 +/- 0.000", "33.33", "0.00"],
    ["SEED_IV",   "33.75 +/- 11.91", "33.33 +/- 0.00", "0.000 +/- 0.000", "33.33", "0.00"],
    ["FACED",     "30.40 +/- 10.84", "33.33 +/- 0.00", "0.000 +/- 0.000", "33.33", "0.00"],
    ["ds004572",  "41.52 +/- 1.17",  "38.67 +/- 1.21", "0.081 +/- 0.019", "33.44", "+5.23"],
    ["ds006437",  "80.52 +/- 0.32",  "68.21 +/- 0.52", "0.553 +/- 0.011", "33.51", "+34.70"],
    ["Overall",   "43.18 +/- 13.40", "38.66 +/- 11.21", "0.082 +/- 0.180", "33.64", "+5.02"],
]
tbl = doc.add_table(rows=1, cols=len(cols))
try:
    tbl.style = "Table Grid"
except Exception:
    pass
hdr = tbl.rows[0].cells
for i, c in enumerate(cols):
    set_run_text(hdr[i].paragraphs[0], c, bold=True)
for r in rows:
    cells = tbl.add_row().cells
    for i, v in enumerate(r):
        set_run_text(cells[i].paragraphs[0], v)

cap = new_para()
set_run_text(cap,
    "Table 3. Within-domain upper bound (same 63-dim features, RF 200 trees, participant-disjoint "
    "GroupShuffleSplit 80/20, 20 seeds; values mean +/- SD). Cross-domain zero-shot BAcc is reproduced "
    "from Table 4 for direct comparison. Delta = within BAcc minus cross-domain ZS BAcc. Only ds006437 "
    "exceeds the 40% learnable-signal threshold; ds004572 is marginal; the remaining six datasets sit at "
    "or below three-class chance even within-dataset, indicating their proxy labels carry little learnable "
    "class signal.")

# ---------- 5. Insert (head, intro, cap, table) before anchor ----------
anchor_el = anchor._p
anchor_el.addprevious(tbl._tbl)   # [table][anchor]
anchor_el.addprevious(cap)        # [cap][table][anchor]
anchor_el.addprevious(intro)      # [intro][cap][table][anchor]
anchor_el.addprevious(head)       # [head][intro][cap][table][anchor]

doc.save(OUT)
print("SAVED", OUT)
# quick verification
d2 = docx.Document(OUT)
print("tables:", len(d2.tables))
# print 4.x headings
for p in d2.paragraphs:
    if re.match(r'^4\.\d+ ', p.text):
        print("HEAD:", p.text)
# print table captions containing 'Table 3' or 'Table 4'
for p in d2.paragraphs:
    if re.match(r'^Table (3|4|13)\. ', p.text.strip()):
        print("CAP:", p.text[:80])
