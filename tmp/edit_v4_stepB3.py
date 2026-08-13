"""Step B3: elevate MAHNOB recovery + split-leakage quantification to the front of
the contributions list (advisor Sec.10 reframing)."""
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = r"E:/universal_bci_hypnosis/paper_en_submission_v5.docx"
doc = docx.Document(PATH)

def set_run_text(p, text, bold=False):
    if hasattr(p, '_p'):
        p = p._p
    for r in p.findall(qn('w:r')):
        p.remove(r)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        b = OxmlElement('w:b'); rPr.append(b)
    r.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t); p.append(r)

# find first contribution bullet
first = None
for p in doc.paragraphs:
    if p.text.strip().startswith("Simultaneous training on seven diverse source domains"):
        first = p
        break
if first is not None:
    # insert new front bullet
    nb = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr'); ps = OxmlElement('w:pStyle'); ps.set(qn('w:val'), 'List Number')
    pPr.append(ps); nb.append(pPr)
    set_run_text(nb, "Recovery of the real 1-9 arousal self-assessment labels for MAHNOB-HCI from session.xml metadata, and the first quantitative estimate of participant- versus trial-level split leakage (SEED_IV accuracy falls from 50.01% under trial-level partitioning to 25.24% under participant-level grouping, a ~25 percentage-point inflation). These are the two assets from this work most likely to be reused by other EEG studies.")
    first._p.addprevious(nb)
    print("front contribution bullet inserted")
else:
    print("first contribution bullet NOT FOUND")

# reword the existing MAHNOB bullet to avoid duplication (focus on participant identity)
for p in doc.paragraphs:
    if p.text.strip().startswith("Recovery of real 1-9 arousal self-assessment labels for MAHNOB-HCI"):
        set_run_text(p, "The recovered MAHNOB labels also restore participant identity, which is what makes participant-level (rather than trial-level) partitioning possible for that dataset and enables the leakage quantification above.")
        print("MAHNOB bullet reworded")
        break

doc.save(PATH)
print("SAVED", PATH)
