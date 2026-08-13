#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Convert PLOS ONE-formatted paper_final.md -> paper_en_submission_v2.docx

Handles: Title, Heading 1/2/3, paragraphs with inline **bold**/*italic*/`code`,
markdown tables (with bold caption above, ⚠️ removed, footnotes), blockquotes,
ordered/unordered lists, and fenced code blocks. Sets 12pt double-spaced body
and adds page numbers in the footer.
"""
import re
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"E:/universal_bci_hypnosis/paper_final.md"
OUT = r"C:/Users/Administrator/Downloads/paper_en_submission_v2.docx"

# ---------- inline formatting ----------
def add_runs(para, text):
    text = clean_text(text)
    pattern = re.compile(r'(\*\*.+?\*\*|\*.+?\*|`.+?`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            para.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith('**'):
            r = para.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('`'):
            r = para.add_run(tok[1:-1]); r.font.name = 'Consolas'
        else:
            r = para.add_run(tok[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])

def clean_cell(s):
    """Remove the ⚠️ colored marker (U+26A0 + optional U+FE0F variation selector).
    PLOS forbids footnotes, so the marker is dropped entirely (the FACED
    artificial-balance caveat is explained in the body text instead)."""
    s = s.replace('\u26a0\ufe0f', '').replace('\u26a0', '')
    # strip any stray variation selectors / zero-width joiners left behind
    s = ''.join(ch for ch in s if ch not in ('\ufe0f', '\u200b', '\u200c', '\u200d'))
    return s.strip()

def clean_text(s):
    """Strip the ⚠️ marker and stray variation selectors from any prose text."""
    s = s.replace('\u26a0\ufe0f', ' a').replace('\u26a0', ' a')
    s = ''.join(ch for ch in s if ch not in ('\ufe0f', '\u200b', '\u200c', '\u200d'))
    return s

# ---------- document setup ----------
doc = Document()
# base style: 12pt, double spaced
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 2.0
normal.paragraph_format.space_after = Pt(6)

# remove default empty paragraph
if doc.paragraphs and doc.paragraphs[0].text == '':
    p = doc.paragraphs[0]._element
    p.getparent().remove(p)

# ---------- PLOS title-page block (first page of manuscript) ----------
AUTHORS = [("Weng Zexiao", False), ("1", True), (", Jung Minpo", False),
           ("1", True), ("*", True)]
AFFILIATION = "1 Department of Computer Engineering, Youngsan University, Yangsan, Republic of Korea"
CORRESP_EMAIL = "E-mail: minpo@ysu.ac.kr (JM)"
ORCIDS = [("Weng Zexiao", "https://orcid.org/0009-0009-8600-8954"),
          ("Jung Minpo", "[ORCID to be added — PLOS ONE requires an ORCID iD for the corresponding author]")]

def add_title_page_block(doc):
    """Insert title-page content (authors, affiliations, corresponding author,
    ORCID) immediately after the title so the manuscript's first page carries
    title + authors + affiliations per PLOS organization rules."""
    # author byline with superscript numbers / * symbol
    bp = doc.add_paragraph()
    bp.paragraph_format.line_spacing = 1.0
    bp.paragraph_format.space_after = Pt(4)
    for txt, sup in AUTHORS:
        r = bp.add_run(txt)
        if sup:
            r.font.superscript = True
    # ORCID iDs
    ol = doc.add_paragraph()
    ol.paragraph_format.line_spacing = 1.0
    ol.paragraph_format.space_after = Pt(2)
    r = ol.add_run("ORCID iDs:"); r.bold = True
    for name, oid in ORCIDS:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(2)
        p.add_run(f"{name}: {oid}")
    # affiliation
    af = doc.add_paragraph()
    af.paragraph_format.line_spacing = 1.0
    af.paragraph_format.space_after = Pt(4)
    af.add_run(AFFILIATION)
    # corresponding author
    ca = doc.add_paragraph()
    ca.paragraph_format.line_spacing = 1.0
    ca.paragraph_format.space_after = Pt(2)
    r = ca.add_run("* Corresponding author"); r.bold = True
    ce = doc.add_paragraph()
    ce.paragraph_format.line_spacing = 1.0
    ce.paragraph_format.space_after = Pt(8)
    ce.add_run(CORRESP_EMAIL)

lines = open(SRC, encoding='utf-8').read().split('\n')
i = 0
n = len(lines)
pending_caption = None   # bold caption paragraph text to place before next table

def is_table_row(line):
    s = line.strip()
    return s.startswith('|') and s.endswith('|') and s.count('|') >= 2

def parse_table_block(start):
    """Parse consecutive table rows starting at `start`. Returns (rows, end_idx)."""
    rows = []
    j = start
    while j < n and is_table_row(lines[j]):
        cells = [c.strip() for c in lines[j].strip().strip('|').split('|')]
        rows.append(cells)
        j += 1
    return rows, j

while i < n:
    line = lines[i]
    stripped = line.strip()

    # ---- blank ----
    if stripped == '':
        i += 1
        continue

    # ---- Title ----
    if stripped.startswith('# ') and not stripped.startswith('## '):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(stripped[2:].strip())
        r.bold = True
        r.font.size = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        add_title_page_block(doc)   # authors + affiliations on the first page
        i += 1
        continue

    # ---- code fence ----
    if stripped.startswith('```'):
        i += 1
        code_lines = []
        while i < n and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        i += 1  # skip closing fence
        cp = doc.add_paragraph()
        cp.paragraph_format.left_indent = Inches(0.3)
        cp.paragraph_format.line_spacing = 1.0
        r = cp.add_run('\n'.join(code_lines))
        r.font.name = 'Consolas'
        r.font.size = Pt(9)
        continue

    # ---- table block ----
    if is_table_row(line):
        # determine caption FIRST (so it is placed above the table)
        caption = pending_caption
        if caption is None:
            k = i - 1
            while k >= 0 and lines[k].strip() == '':
                k -= 1
            if k >= 0:
                m = re.match(r'\*\*(Table [A-Z]?\d+\..+?)\*\*', lines[k].strip())
                if m:
                    caption = m.group(1)
        pending_caption = None
        if caption is not None:
            cap = doc.add_paragraph()
            cap.paragraph_format.space_after = Pt(2)
            r = cap.add_run(clean_text(caption))
            r.bold = True
        # build table
        rows, j = parse_table_block(i)
        # rows[0]=header, rows[1]=alignment(separator), rows[2:]=data
        header = rows[0]
        data = rows[2:] if len(rows) > 2 else []
        t = doc.add_table(rows=1, cols=len(header))
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        # header row
        hdr = t.rows[0].cells
        for c, txt in enumerate(header):
            hdr[c].text = ''
            add_runs(hdr[c].paragraphs[0], clean_text(txt))
            for run in hdr[c].paragraphs[0].runs:
                run.bold = True
        # data rows
        for drow in data:
            cells = t.add_row().cells
            for c, txt in enumerate(drow):
                txt = clean_cell(txt)
                cells[c].text = ''
                add_runs(cells[c].paragraphs[0], txt)
        i = j
        continue

    # ---- heading 2 (### ) ----
    if stripped.startswith('### '):
        p = doc.add_paragraph(clean_text(stripped[4:].strip()), style='Heading 2')
        i += 1
        continue
    # ---- heading 1 (## ) ----
    if stripped.startswith('## '):
        p = doc.add_paragraph(clean_text(stripped[3:].strip()), style='Heading 1')
        i += 1
        continue

    # ---- table caption (**Table N. ...**) -> defer before next table ----
    mcap = re.match(r'\*\*(Table [A-Z]?\d+\..+?)\*\*', stripped)
    if mcap:
        pending_caption = mcap.group(1)
        i += 1
        continue

    # ---- blockquote (> ) ----
    if stripped.startswith('> '):
        txt = clean_text(stripped[2:].strip())
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(txt)
        r.italic = True
        r.font.size = Pt(10)
        i += 1
        continue

    # ---- unordered list (- ) ----
    if re.match(r'^-\s+', stripped):
        p = doc.add_paragraph(style='List Bullet')
        add_runs(p, stripped[2:].strip())
        i += 1
        continue

    # ---- ordered list (N. ) ----
    if re.match(r'^\d+\.\s+', stripped):
        p = doc.add_paragraph(style='List Number')
        add_runs(p, re.sub(r'^\d+\.\s+', '', stripped))
        i += 1
        continue

    # ---- normal paragraph ----
    p = doc.add_paragraph()
    add_runs(p, stripped)
    i += 1

# ---------- page numbers in footer ----------
section = doc.sections[0]
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
# PAGE field must be a direct child of <w:p>, not nested inside a <w:r>
fld1 = OxmlElement('w:fldSimple')
fld1.set(qn('w:instr'), ' PAGE ')
r = OxmlElement('w:r')
t = OxmlElement('w:t')
t.text = '1'  # cached value; Word recalculates on open
r.append(t)
fld1.append(r)
fp._p.append(fld1)

# ---------- continuous line numbers (PLOS requirement) ----------
def add_line_numbers(document, count_by=1, restart='continuous', distance='360'):
    """Enable Word line numbering on every section. PLOS requires continuous
    line numbers (must not restart on each page)."""
    for sec in document.sections:
        sectPr = sec._sectPr
        ln = sectPr.find(qn('w:lnNumType'))
        if ln is None:
            ln = OxmlElement('w:lnNumType')
            sectPr.append(ln)
        ln.set(qn('w:countBy'), str(count_by))
        ln.set(qn('w:restart'), restart)
        ln.set(qn('w:distance'), distance)

add_line_numbers(doc)

doc.save(OUT)
print(f"Saved {OUT}")
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
