"""Step B4: fix contribution-list styling and MAHNOB bullet wording."""
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = r"E:/universal_bci_hypnosis/paper_en_submission_v5.docx"
doc = docx.Document(PATH)

def set_run_text(p, text):
    if hasattr(p, '_p'):
        p = p._p
    for r in p.findall(qn('w:r')):
        p.remove(r)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t); p.append(r)

# 1) new front bullet -> List Number style
for p in doc.paragraphs:
    if p.text.strip().startswith("Recovery of the real 1-9 arousal self-assessment labels for MAHNOB-HCI"):
        p.style = doc.styles['List Number']
        print("front bullet styled:", p.style.name)
        break

# 2) reword the ORIGINAL MAHNOB bullet (starts with 'Recovery of real 1-9 ... which also restores')
for p in doc.paragraphs:
    if p.text.strip().startswith("Recovery of real 1-9 arousal self-assessment labels for MAHNOB-HCI from session.xml metadata, which"):
        set_run_text(p, "The recovered MAHNOB labels also restore participant identity, which is what makes participant-level (rather than trial-level) partitioning possible for that dataset and enables the leakage quantification listed above.")
        print("MAHNOB bullet reworded")
        break

doc.save(PATH)
print("SAVED", PATH)
