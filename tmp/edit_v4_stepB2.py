"""Step B2: fix footnote insertion (C2) and C4 funding/competing interests."""
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = r"E:/universal_bci_hypnosis/paper_en_submission_v5.docx"
doc = docx.Document(PATH)

def set_run_text(p, text, bold=False):
    if hasattr(p, '_p'):
        p = p._p
    for r in p.findall(qn('w:r')):
        p.remove(r)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        b = OxmlElement('w:b'); rPr.append(b)
    r.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t); p.append(r)

def new_para():
    return OxmlElement('w:p')

# ---- C2 footnote after the "+ Only 10 ..." paragraph ----
foot = None
for p in doc.paragraphs:
    if 'Only 10 of the 15 publicly documented SEED' in p.text:
        foot = p
if foot is not None:
    fp = new_para()
    set_run_text(fp, "++ ds006437 comprises 9 participants across 36 sessions; it is partitioned by session identifier, so within-participant leakage cannot be fully excluded and its within-domain estimate (Table 3) uses session-level grouping.")
    foot._p.addnext(fp)
    print("footnote inserted")
else:
    print("footnote anchor NOT FOUND")

# ---- C4 after the "A completed checklist mapping each STROBE..." paragraph ----
body73 = None
for p in doc.paragraphs:
    if 'A completed checklist mapping each STROBE' in p.text:
        body73 = p
if body73 is not None:
    p1 = new_para(); set_run_text(p1, "The authors received no specific funding for this work.")
    p2 = new_para(); set_run_text(p2, "The authors have declared that no competing interests exist.")
    body73._p.addnext(p2)
    body73._p.addnext(p1)
    print("C4 inserted")
else:
    print("C4 anchor NOT FOUND")

doc.save(PATH)
print("SAVED", PATH)
