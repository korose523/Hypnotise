import docx, re, shutil, os

V6 = r"E:/universal_bci_hypnosis/paper_en_submission_v6.docx"
V5 = r"E:/universal_bci_hypnosis/paper_en_submission_v5.docx"
BAK = r"E:/universal_bci_hypnosis/paper_en_submission_v6_backup_before_v1.0.13.docx"
if not os.path.exists(BAK):
    shutil.copyfile(V6, BAK)
    print("backed up v6 ->", BAK)

OLD_VER_REC = "v1.0.12 versioned record https://doi.org/10.5281/zenodo.21922749"
NEW_VER_REC = "v1.0.13 versioned record https://doi.org/10.5281/zenodo.21922961"

def count_in(doc, pat):
    n = 0
    for p in doc.paragraphs:
        n += len(re.findall(pat, p.text))
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                n += len(re.findall(pat, c.text))
    return n

def replace_in_doc(doc, old, new):
    made = 0
    def do_runs(runs):
        nonlocal made
        full = "".join(r.text for r in runs)
        if old not in full:
            return
        new_full = full.replace(old, new)
        runs[0].text = new_full
        for r in runs[1:]:
            r.text = ""
        made += full.count(old)
    for p in doc.paragraphs:
        do_runs(p.runs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for para in c.paragraphs:
                    do_runs(para.runs)
    return made

d = docx.Document(V6)
print("PRE  versioned-rec:", count_in(d, re.escape(OLD_VER_REC)),
      "| tag v1.0.12:", count_in(d, r"tag v1\.0\.12"),
      "| 21922749:", count_in(d, r"21922749"))

m1 = replace_in_doc(d, OLD_VER_REC, NEW_VER_REC)
m2 = replace_in_doc(d, "tag v1.0.12", "tag v1.0.13")

print("Made versioned-rec:", m1, "| tag:", m2)
print("POST versioned-rec:", count_in(d, re.escape(NEW_VER_REC)),
      "| 21922961:", count_in(d, r"21922961"),
      "| leftover v1.0.12:", count_in(d, r"v1\.0\.12"))
d.save(V6)
print("saved v6:", os.path.getsize(V6))
shutil.copyfile(V6, V5)
print("synced v5:", os.path.getsize(V5))

def edit_file(path, repls):
    s = open(path, encoding="utf-8").read()
    for old, new in repls:
        c = s.count(old)
        s = s.replace(old, new)
        print("  %s: %r x%d" % (os.path.basename(path), old[:55], c))
    open(path, "w", encoding="utf-8").write(s)

print("=== md edits ===")
edit_file(r"E:/universal_bci_hypnosis/plos_human_participants_checklist.md", [
    (OLD_VER_REC, NEW_VER_REC),
    ("tag v1.0.12", "tag v1.0.13"),
])
edit_file(r"E:/universal_bci_hypnosis/em_compliance_answers.md", [
    ("v1.0.12", "v1.0.13"),
    ("21922749", "21922961"),
])
edit_file(r"E:/universal_bci_hypnosis/paper_final.md", [
    ("zenodo.21531272` [v1.0.12", "zenodo.21922961` [v1.0.13"),
])
edit_file(r"E:/universal_bci_hypnosis/README.md", [
    ("tag v1.0.12", "tag v1.0.13"),
    ("record **21920943**", "record **21922961**"),
])
print("DONE")
